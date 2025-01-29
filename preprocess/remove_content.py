from docx import Document


def delete_sections(document_path, output_path):
    
    doc = Document(document_path)
    
    titles_to_delete = ["References"]

    def remove_paragraph(paragraph):
        
        p = paragraph._element
        p.getparent().remove(p)
        p._element = p._p = None

    def remove_content_after_heading(doc, heading_text, delete_to_next_heading=True):
        
        delete = False
        for paragraph in doc.paragraphs:

            if heading_text in paragraph.text and paragraph.style.name.startswith("Heading"):
                delete = True
                remove_paragraph(paragraph)
                continue

            if delete:
                if delete_to_next_heading and paragraph.style.name.startswith("Heading"):
                    break
                remove_paragraph(paragraph)

    delete_contents = False
    for paragraph in doc.paragraphs:
        if "Contents" in paragraph.text:
            delete_contents = True
        if delete_contents:
            if paragraph.style.name.startswith("Heading") and "Contents" not in paragraph.text:
                break 
            remove_paragraph(paragraph)

    for title in titles_to_delete:
        remove_content_after_heading(doc, title)

    found_annex = False
    paragraphs_to_delete = []
    for i, paragraph in enumerate(doc.paragraphs):
        if "Annex" in paragraph.text and paragraph.style.name.startswith("Heading"):
            found_annex = True
        if found_annex:
            paragraphs_to_delete.append(paragraph)

    for paragraph in paragraphs_to_delete:
        remove_paragraph(paragraph)

    for table in doc.tables:
        tbl = table._element
        tbl.getparent().remove(tbl)

    doc.save(output_path)